import io
import re

from docx import Document
from docx.shared import Pt
from fpdf import FPDF


def strip_markdown(text: str) -> str:
    """Remove markdown formatting from text."""
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove heading markers
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    # Remove horizontal rules
    text = re.sub(r'^-{3,}$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^\*{3,}$', '', text, flags=re.MULTILINE)
    return text


def is_title_line(line: str) -> bool:
    """Check if a line is a title (mostly uppercase, short)."""
    stripped = line.strip()
    if not stripped or len(stripped) < 3:
        return False
    # Count uppercase letters
    upper_count = sum(1 for c in stripped if c.isupper())
    letter_count = sum(1 for c in stripped if c.isalpha())
    if letter_count == 0:
        return False
    # Consider it a title if >70% uppercase and reasonably short
    return (upper_count / letter_count > 0.7) and len(stripped) < 80


def text_to_docx(text: str, title: str = "Documento") -> bytes:
    """Convert plain text to DOCX format."""
    text = strip_markdown(text)
    doc = Document()

    lines = text.split("\n")
    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Empty line - add paragraph break
            doc.add_paragraph()
        elif is_title_line(stripped):
            # Title/heading line
            para = doc.add_paragraph()
            run = para.add_run(stripped)
            run.bold = True
            run.font.size = Pt(12)
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        else:
            # Normal paragraph
            para = doc.add_paragraph(line)
            para.paragraph_format.space_after = Pt(3)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_margins(20, 20, 20)

    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Pagina {self.page_no()}", 0, 0, "C")


def sanitize_text(text: str) -> str:
    """Remove or replace characters that cause issues in PDF."""
    replacements = {
        '\u2018': "'",  # Left single quote
        '\u2019': "'",  # Right single quote
        '\u201c': '"',  # Left double quote
        '\u201d': '"',  # Right double quote
        '\u2013': '-',  # En dash
        '\u2014': '-',  # Em dash
        '\u2026': '...',  # Ellipsis
        '\u00a0': ' ',  # Non-breaking space
        '\u00ba': 'o',  # Masculine ordinal
        '\u00aa': 'a',  # Feminine ordinal
        '\u00b0': 'o',  # Degree symbol
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def text_to_pdf(text: str, title: str = "Documento") -> bytes:
    """Convert plain text to PDF format."""
    text = strip_markdown(text)
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=25)

    effective_width = pdf.w - pdf.l_margin - pdf.r_margin

    lines = text.split("\n")
    for line in lines:
        line_sanitized = sanitize_text(line)
        stripped = line_sanitized.strip()

        if not stripped:
            pdf.ln(4)
        elif is_title_line(line.strip()):  # Check original for case
            pdf.set_font("Helvetica", "B", 11)
            pdf.ln(3)
            pdf.multi_cell(effective_width, 6, stripped)
            pdf.set_font("Helvetica", "", 10)
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(effective_width, 5, line_sanitized)

    return bytes(pdf.output())


def format_document(text: str, format_type: str, title: str = "Documento") -> tuple:
    """
    Format document to specified type.
    Returns (bytes, mime_type, file_extension)
    """
    # Strip markdown from text first
    clean_text = strip_markdown(text)

    if format_type == "docx":
        return (
            text_to_docx(clean_text, title),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".docx"
        )
    elif format_type == "pdf":
        return (
            text_to_pdf(clean_text, title),
            "application/pdf",
            ".pdf"
        )
    else:  # txt
        return (
            clean_text.encode("utf-8"),
            "text/plain",
            ".txt"
        )
